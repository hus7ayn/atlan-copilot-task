#!/usr/bin/env python3
"""
File Parser for Various Document Formats
Supports PDF, DOCX, TXT, CSV, JSON, and more
"""

import os
import json
import csv
import io
import hashlib
from typing import Dict, List, Any, Optional, Union
from pathlib import Path
from datetime import datetime
import PyPDF2
from docx import Document
# import pandas as pd  # Removed to avoid compilation issues
from bs4 import BeautifulSoup
import re

class FileParser:
    """Universal file parser for various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,  # Will try docx parser
            '.txt': self._parse_txt,
            '.csv': self._parse_csv,
            '.json': self._parse_json,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.md': self._parse_markdown,
            '.log': self._parse_log,
            '.xml': self._parse_xml,
            '.yaml': self._parse_yaml,
            '.yml': self._parse_yaml
        }
    
    def parse_file(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """
        Parse a file and extract its content
        
        Args:
            file_path: Path to the file or filename for content
            file_content: Raw file content as bytes (optional)
        
        Returns:
            Dictionary with parsed content and metadata
        """
        try:
            # Get file extension
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_ext}',
                    'supported_formats': list(self.supported_formats.keys())
                }
            
            # Parse the file
            if file_content:
                # Parse from bytes
                content = self.supported_formats[file_ext](file_content)
            else:
                # Parse from file path
                with open(file_path, 'rb') as f:
                    content = self.supported_formats[file_ext](f.read())
            
            return {
                'success': True,
                'content': content,
                'file_type': file_ext,
                'file_name': Path(file_path).name,
                'word_count': len(content.split()) if isinstance(content, str) else 0,
                'char_count': len(content) if isinstance(content, str) else 0
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error parsing file: {str(e)}',
                'file_type': file_ext if 'file_ext' in locals() else 'unknown'
            }
    
    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF content"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"PDF parsing error: {str(e)}")
    
    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX content"""
        try:
            doc = Document(io.BytesIO(content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"DOCX parsing error: {str(e)}")
    
    def _parse_txt(self, content: bytes) -> str:
        """Parse plain text content"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return content.decode('utf-8', errors='replace')
        except Exception as e:
            raise Exception(f"Text parsing error: {str(e)}")
    
    def _parse_csv(self, content: bytes) -> str:
        """Parse CSV content"""
        try:
            text_content = content.decode('utf-8')
            csv_reader = csv.reader(io.StringIO(text_content))
            
            # Convert CSV to readable text
            lines = []
            for row in csv_reader:
                lines.append(" | ".join(row))
            
            return "\n".join(lines)
        except Exception as e:
            raise Exception(f"CSV parsing error: {str(e)}")
    
    def _parse_json(self, content: bytes) -> str:
        """Parse JSON content"""
        try:
            json_data = json.loads(content.decode('utf-8'))
            
            # Convert JSON to readable text
            if isinstance(json_data, dict):
                return self._dict_to_text(json_data)
            elif isinstance(json_data, list):
                return self._list_to_text(json_data)
            else:
                return str(json_data)
        except Exception as e:
            raise Exception(f"JSON parsing error: {str(e)}")
    
    def _parse_html(self, content: bytes) -> str:
        """Parse HTML content"""
        try:
            soup = BeautifulSoup(content.decode('utf-8'), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            raise Exception(f"HTML parsing error: {str(e)}")
    
    def _parse_markdown(self, content: bytes) -> str:
        """Parse Markdown content"""
        try:
            text = content.decode('utf-8')
            # Basic markdown cleaning - remove markdown syntax
            text = re.sub(r'#{1,6}\s+', '', text)  # Remove headers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Remove code
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links
            return text
        except Exception as e:
            raise Exception(f"Markdown parsing error: {str(e)}")
    
    def _parse_log(self, content: bytes) -> str:
        """Parse log file content"""
        try:
            return content.decode('utf-8')
        except Exception as e:
            raise Exception(f"Log parsing error: {str(e)}")
    
    def _parse_xml(self, content: bytes) -> str:
        """Parse XML content"""
        try:
            soup = BeautifulSoup(content.decode('utf-8'), 'xml')
            return soup.get_text()
        except Exception as e:
            raise Exception(f"XML parsing error: {str(e)}")
    
    def _parse_yaml(self, content: bytes) -> str:
        """Parse YAML content"""
        try:
            import yaml
            yaml_data = yaml.safe_load(content.decode('utf-8'))
            
            if isinstance(yaml_data, dict):
                return self._dict_to_text(yaml_data)
            elif isinstance(yaml_data, list):
                return self._list_to_text(yaml_data)
            else:
                return str(yaml_data)
        except Exception as e:
            raise Exception(f"YAML parsing error: {str(e)}")
    
    def _dict_to_text(self, data: dict, indent: int = 0) -> str:
        """Convert dictionary to readable text"""
        lines = []
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append("  " * indent + f"{key}:")
                lines.append(self._dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                lines.append("  " * indent + f"{key}:")
                lines.append(self._list_to_text(value, indent + 1))
            else:
                lines.append("  " * indent + f"{key}: {value}")
        return "\n".join(lines)
    
    def _list_to_text(self, data: list, indent: int = 0) -> str:
        """Convert list to readable text"""
        lines = []
        for i, item in enumerate(data):
            if isinstance(item, dict):
                lines.append("  " * indent + f"Item {i + 1}:")
                lines.append(self._dict_to_text(item, indent + 1))
            elif isinstance(item, list):
                lines.append("  " * indent + f"Item {i + 1}:")
                lines.append(self._list_to_text(item, indent + 1))
            else:
                lines.append("  " * indent + f"Item {i + 1}: {item}")
        return "\n".join(lines)
    
    def extract_tickets_from_content(self, content: str) -> List[Dict[str, Any]]:
        """
        Extract potential tickets from parsed content with flexible parsing
        
        Args:
            content: Parsed text content
        
        Returns:
            List of potential tickets
        """
        tickets = []
        
        # Clean and normalize the content
        content = self._normalize_content(content)
        
        # Try multiple extraction strategies
        extraction_methods = [
            self._extract_by_ticket_patterns,
            self._extract_by_paragraphs,
            self._extract_by_sections,
            self._extract_by_bullets,
            self._extract_by_emails,
            self._extract_by_json_format,
            self._extract_by_structured_format
        ]
        
        for method in extraction_methods:
            try:
                extracted_tickets = method(content)
                if extracted_tickets:
                    tickets.extend(extracted_tickets)
                    break  # Use first successful method
            except Exception as e:
                print(f"⚠️ Extraction method {method.__name__} failed: {e}")
                continue
        
        # If no tickets found, treat entire content as one ticket
        if not tickets and content.strip():
            tickets.append(self._create_ticket_from_content(content, "FULL-CONTENT-1"))
        
        # Post-process and validate tickets
        tickets = self._post_process_tickets(tickets)
        
        return tickets
    
    def _normalize_content(self, content: str) -> str:
        """Normalize content for better parsing"""
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)
        # Normalize line endings
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        # Remove leading/trailing whitespace
        content = content.strip()
        return content
    
    def _extract_by_ticket_patterns(self, content: str) -> List[Dict[str, Any]]:
        """Extract tickets using common ticket patterns"""
        tickets = []
        
        # Enhanced patterns for ticket identification
        patterns = [
            r'(?:Ticket|Issue|Case|Request|Support|Bug|Task)\s*#?\s*(\d+)[:\s]*(.*?)(?=(?:Ticket|Issue|Case|Request|Support|Bug|Task)\s*#?\s*\d+|\Z)',
            r'(?:ID|Id|ID:)\s*(\d+)[:\s]*(.*?)(?=(?:ID|Id|ID:)\s*\d+|\Z)',
            r'(\d+)[\.\)]\s*(.*?)(?=\d+[\.\)]\s|\Z)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, content, re.DOTALL | re.IGNORECASE)
            for match in matches:
                ticket_id = match.group(1) if len(match.groups()) >= 1 else f"TICKET-{len(tickets) + 1}"
                ticket_content = match.group(2) if len(match.groups()) >= 2 else match.group(0)
                
                if len(ticket_content.strip()) > 10:
                    tickets.append(self._create_ticket_from_content(ticket_content, ticket_id))
        
        return tickets
    
    def _extract_by_paragraphs(self, content: str) -> List[Dict[str, Any]]:
        """Extract tickets by splitting into meaningful paragraphs"""
        tickets = []
        paragraphs = re.split(r'\n\s*\n', content)
        
        for i, paragraph in enumerate(paragraphs):
            paragraph = paragraph.strip()
            if len(paragraph) > 20:  # Minimum meaningful length
                tickets.append(self._create_ticket_from_content(paragraph, f"PARAGRAPH-{i + 1}"))
        
        return tickets
    
    def _extract_by_sections(self, content: str) -> List[Dict[str, Any]]:
        """Extract tickets by identifying sections (headers, etc.)"""
        tickets = []
        
        # Split by headers (lines starting with #, *, -, numbers, etc.)
        sections = re.split(r'\n(?=#{1,6}\s+|\*\s+|\-\s+|\d+\.\s+|[A-Z][^a-z]*:)', content)
        
        for i, section in enumerate(sections):
            section = section.strip()
            if len(section) > 15:
                tickets.append(self._create_ticket_from_content(section, f"SECTION-{i + 1}"))
        
        return tickets
    
    def _extract_by_bullets(self, content: str) -> List[Dict[str, Any]]:
        """Extract tickets from bulleted lists"""
        tickets = []
        
        # Find bulleted items
        bullet_pattern = r'(?:^|\n)\s*[•\*\-\+]\s+(.+?)(?=\n\s*[•\*\-\+]\s+|\n\n|\Z)'
        matches = re.finditer(bullet_pattern, content, re.DOTALL)
        
        for i, match in enumerate(matches):
            bullet_content = match.group(1).strip()
            if len(bullet_content) > 10:
                tickets.append(self._create_ticket_from_content(bullet_content, f"BULLET-{i + 1}"))
        
        return tickets
    
    def _extract_by_emails(self, content: str) -> List[Dict[str, Any]]:
        """Extract tickets from email-like content"""
        tickets = []
        
        # Look for email patterns
        email_pattern = r'(?:From|To|Subject|Date):\s*(.*?)(?=(?:From|To|Subject|Date):|\Z)'
        sections = re.split(email_pattern, content, flags=re.IGNORECASE)
        
        current_email = {}
        for section in sections:
            if section.strip():
                if ':' in section and any(keyword in section.lower() for keyword in ['from', 'to', 'subject', 'date']):
                    # This is a header
                    key, value = section.split(':', 1)
                    current_email[key.strip().lower()] = value.strip()
                else:
                    # This is email body
                    if len(section.strip()) > 10:
                        tickets.append(self._create_ticket_from_content(section.strip(), f"EMAIL-{len(tickets) + 1}"))
                        current_email = {}
        
        return tickets
    
    def _extract_by_json_format(self, content: str) -> List[Dict[str, Any]]:
        """Extract tickets from JSON-like content"""
        tickets = []
        
        try:
            # Try to parse as JSON
            if content.strip().startswith('{') or content.strip().startswith('['):
                json_data = json.loads(content)
                
                if isinstance(json_data, list):
                    for i, item in enumerate(json_data):
                        if isinstance(item, dict):
                            ticket_content = self._dict_to_text(item)
                            tickets.append(self._create_ticket_from_content(ticket_content, f"JSON-{i + 1}"))
                elif isinstance(json_data, dict):
                    ticket_content = self._dict_to_text(json_data)
                    tickets.append(self._create_ticket_from_content(ticket_content, "JSON-1"))
        except:
            # Not JSON, continue with other methods
            pass
        
        return tickets
    
    def _extract_by_structured_format(self, content: str) -> List[Dict[str, Any]]:
        """Extract tickets from structured formats (CSV-like, tab-separated, etc.)"""
        tickets = []
        
        # Try CSV-like format
        lines = content.split('\n')
        if len(lines) > 1:
            # Check if it looks like CSV
            first_line = lines[0]
            if ',' in first_line or '\t' in first_line:
                # Skip header line
                for i, line in enumerate(lines[1:], 1):
                    line = line.strip()
                    if line:
                        tickets.append(self._create_ticket_from_content(line, f"ROW-{i}"))
        
        return tickets
    
    def _create_ticket_from_content(self, content: str, ticket_id: str) -> Dict[str, Any]:
        """Create a ticket object from content"""
        # Extract subject (first line or first sentence)
        lines = content.split('\n')
        first_line = lines[0].strip()
        
        # Use first line as subject, or first sentence if line is too long
        if len(first_line) > 100:
            # Find first sentence
            sentences = re.split(r'[.!?]+', first_line)
            subject = sentences[0].strip() if sentences else first_line[:100]
        else:
            subject = first_line
        
        return {
            'id': ticket_id,
            'subject': subject[:100] if subject else 'Extracted Content',
            'body': content,
            'created_at': datetime.now().isoformat(),
            'customer_email': 'extracted@example.com'
        }
    
    def _post_process_tickets(self, tickets: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Post-process and validate extracted tickets"""
        processed_tickets = []
        
        for ticket in tickets:
            # Clean up the ticket
            ticket['subject'] = self._clean_text(ticket['subject'])
            ticket['body'] = self._clean_text(ticket['body'])
            
            # Validate minimum content
            if len(ticket['body'].strip()) >= 10:
                processed_tickets.append(ticket)
        
        # Remove duplicates based on content similarity
        processed_tickets = self._remove_duplicate_tickets(processed_tickets)
        
        return processed_tickets
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\.\,\!\?\:\;\-\(\)\[\]\{\}\"\'\/\\]', '', text)
        return text.strip()
    
    def _remove_duplicate_tickets(self, tickets: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Remove duplicate tickets based on content similarity"""
        unique_tickets = []
        seen_content = set()
        
        for ticket in tickets:
            # Create a simple hash of the content for comparison
            content_hash = hashlib.md5(ticket['body'].encode()).hexdigest()
            
            if content_hash not in seen_content:
                seen_content.add(content_hash)
                unique_tickets.append(ticket)
        
        return unique_tickets

# Global parser instance
file_parser = FileParser()
